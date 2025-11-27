"""
문서 텍스트 추출 모듈
"""
import os
import zipfile
from pathlib import Path
from typing import Dict, Callable
import PyPDF2
import docx
from utils.constants import MAX_FILE_SIZE, SUPPORTED_EXTENSIONS
from utils.logger import logger


class DocumentProcessor:
    """문서 텍스트 추출"""
    
    def __init__(self):
        self.extractors: Dict[str, Callable] = {
            '.pdf': self._extract_from_pdf,
            '.docx': self._extract_from_docx,
            '.txt': self._extract_from_txt,
            '.hwp': self._extract_from_hwp,
            '.hwpx': self._extract_from_hwpx,
        }
    
    def check_file_size(self, file_path: str) -> bool:
        """파일 크기 확인"""
        size = os.path.getsize(file_path)
        return size <= MAX_FILE_SIZE
    
    def extract_text(self, file_path: str) -> str:
        """
        파일에서 텍스트 추출
        
        Args:
            file_path: 파일 경로
            
        Returns:
            추출된 텍스트
            
        Raises:
            Exception: 파일 처리 중 오류 발생
        """
        if not os.path.exists(file_path):
            raise Exception("파일을 찾을 수 없습니다.")
        
        if not self.check_file_size(file_path):
            raise Exception(f"파일이 너무 큽니다. (최대 {MAX_FILE_SIZE // (1024*1024)}MB)")
        
        ext = Path(file_path).suffix.lower()
        
        if ext not in SUPPORTED_EXTENSIONS:
            raise Exception(f"지원하지 않는 파일 형식: {ext}")
        
        extractor = self.extractors.get(ext)
        if not extractor:
            raise Exception(f"추출기를 찾을 수 없습니다: {ext}")
        
        return extractor(file_path)
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """PDF 텍스트 추출"""
        try:
            text = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                max_pages = min(total_pages, 100)
                
                for page in pdf_reader.pages[:max_pages]:
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text.append(page_text)
                    except:
                        continue
            
            result = "\n".join(text)
            if not result.strip():
                raise Exception("PDF에서 텍스트를 추출할 수 없습니다.")
            return result
        except Exception as e:
            raise Exception(f"PDF 처리 오류: {str(e)}")
    
    def _extract_from_docx(self, file_path: str) -> str:
        """DOCX 텍스트 추출"""
        try:
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # 테이블 텍스트도 추출
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            table_texts.append(cell.text.strip())
            
            return "\n".join(paragraphs + table_texts)
        except Exception as e:
            raise Exception(f"DOCX 처리 오류: {str(e)}")
    
    def _extract_from_txt(self, file_path: str) -> str:
        """TXT 텍스트 추출"""
        encodings = ['utf-8', 'cp949', 'euc-kr', 'latin-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except:
                continue
        
        raise Exception("TXT 파일 인코딩을 인식할 수 없습니다.")
    
    def _extract_from_hwp(self, file_path: str) -> str:
        """HWP 텍스트 추출 (pyhwp 기반)"""
        try:
            # pyhwp 임포트 확인
            import hwp5
            from hwp5.binmodel import Hwp5File
            
            logger.info(f"pyhwp 버전: {hwp5.__version__ if hasattr(hwp5, '__version__') else '확인불가'}")
            
        except ImportError as ie:
            logger.error(f"pyhwp 라이브러리 임포트 오류: {str(ie)}")
            raise Exception("pyhwp 라이브러리가 필요합니다. 'pip install pyhwp' 명령으로 설치하세요.")
        
        # 텍스트 추출 시도 - 가장 안정적인 방법부터
        try:
            # 방법 1: subprocess로 hwp5txt 명령 실행 (가장 안정적)
            try:
                import subprocess
                import tempfile
                import os
                
                # hwp5txt 명령이 사용 가능한지 확인
                try:
                    subprocess.run(['hwp5txt', '--help'], capture_output=True, timeout=5)
                except:
                    raise Exception("hwp5txt 명령을 찾을 수 없습니다")
                
                # 임시 출력 파일 생성
                with tempfile.NamedTemporaryFile(mode='w+', suffix='.txt', delete=False, encoding='utf-8') as tmp:
                    tmp_path = tmp.name
                
                # hwp5txt 명령 실행
                result = subprocess.run(
                    ['hwp5txt', file_path],
                    capture_output=True,
                    text=True,
                    timeout=30,
                    encoding='utf-8'
                )
                
                if result.returncode == 0 and result.stdout:
                    extracted_text = result.stdout.strip()
                    if extracted_text:
                        logger.info(f"HWP 파일 추출 완료 (방법1-stdout): {len(extracted_text)} 문자")
                        return extracted_text
                
                # stdout이 없으면 파일 출력 시도
                result = subprocess.run(
                    ['hwp5txt', file_path, '-o', tmp_path],
                    capture_output=True,
                    text=True,
                    timeout=30,
                    encoding='utf-8'
                )
                
                # 결과 읽기
                if os.path.exists(tmp_path):
                    with open(tmp_path, 'r', encoding='utf-8') as f:
                        extracted_text = f.read()
                    
                    # 임시 파일 삭제
                    os.unlink(tmp_path)
                    
                    if extracted_text and extracted_text.strip():
                        logger.info(f"HWP 파일 추출 완료 (방법1-file): {len(extracted_text)} 문자")
                        return extracted_text
            
            except Exception as e1:
                logger.warning(f"방법1 실패: {str(e1)}")
            
            # 방법 2: Python API 사용 - 가장 간단한 접근
            try:
                hwp = Hwp5File(file_path)
                
                # preview_text 시도
                if hasattr(hwp, 'preview_text') and hwp.preview_text:
                    try:
                        result = str(hwp.preview_text).strip()
                        if result and len(result) > 10:
                            hwp.close()
                            logger.info(f"HWP 파일 추출 완료 (방법2-preview): {len(result)} 문자")
                            return result
                    except:
                        pass
                
                # bodytext 직접 접근 시도
                if hasattr(hwp, 'bodytext') and hwp.bodytext:
                    try:
                        text_parts = []
                        bodytext = hwp.bodytext
                        
                        # bodytext를 문자열로 변환 시도
                        bodytext_str = str(bodytext)
                        if bodytext_str and len(bodytext_str) > 10:
                            text_parts.append(bodytext_str)
                        
                        if text_parts:
                            result = '\n'.join(text_parts)
                            hwp.close()
                            logger.info(f"HWP 파일 추출 완료 (방법2-bodytext): {len(result)} 문자")
                            return result
                    except:
                        pass
                
                hwp.close()
            
            except Exception as e2:
                logger.warning(f"방법2 실패: {str(e2)}")
            
            # 방법 3: hwp5txt 모듈 직접 사용
            try:
                from hwp5.hwp5txt import TextTransform
                from io import StringIO
                
                hwp = Hwp5File(file_path)
                transform = TextTransform()
                
                # StringIO를 사용하여 텍스트 출력 캡처
                output = StringIO()
                transform.transform(hwp, output)
                result = output.getvalue()
                
                hwp.close()
                output.close()
                
                if result and result.strip():
                    logger.info(f"HWP 파일 추출 완료 (방법3): {len(result)} 문자")
                    return result
            
            except Exception as e3:
                logger.warning(f"방법3 실패: {str(e3)}")
            
            # 모든 방법 실패
            raise Exception("HWP 파일에서 텍스트를 추출할 수 없습니다. 파일이 손상되었거나 암호화되었을 수 있습니다.")
        
        except Exception as e:
            logger.error(f"HWP 처리 오류: {str(e)}")
            raise Exception(f"HWP 처리 오류: {str(e)}")
    
    def _extract_from_hwpx(self, file_path: str) -> str:
        """HWPX 텍스트 추출"""
        try:
            import xml.etree.ElementTree as ET
            text_parts = []
            
            with zipfile.ZipFile(file_path, 'r') as zf:
                sections = sorted([
                    f for f in zf.namelist() 
                    if f.startswith('Contents/section') and f.endswith('.xml')
                ])
                
                for section in sections:
                    try:
                        content = zf.open(section).read().decode('utf-8')
                        root = ET.fromstring(content)
                        texts = [elem.text for elem in root.iter() if elem.text]
                        text_parts.append(' '.join(texts))
                    except:
                        continue
            
            return '\n\n'.join(text_parts)
        except Exception as e:
            raise Exception(f"HWPX 처리 오류: {str(e)}")
